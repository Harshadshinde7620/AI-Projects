import sys
import json
from docx import Document
import os

def generate_test_plan(data, template_path, output_path):
    """
    Populates a .docx template with test plan data.
    """
    if not os.path.exists(template_path):
        return {"status": "error", "message": f"Template not found at {template_path}"}
    
    try:
        doc = Document(template_path)
        
        # Simple placeholder replacement
        placeholders = {
            "{{JIRA_ID}}": data.get("jira_id", "N/A"),
            "{{SUMMARY}}": data.get("summary", "N/A"),
            "{{DESCRIPTION}}": data.get("description", "N/A"),
            "{{TEST_CASES}}": data.get("test_cases", "N/A"),
            "{{PROJECT}}": data.get("project", "N/A")
        }
        
        for paragraph in doc.paragraphs:
            for key, value in placeholders.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
        
        doc.save(output_path)
        return {"status": "success", "file_path": output_path}
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"status": "error", "message": "No input payload"}))
        sys.exit(1)
        
    try:
        payload = json.loads(sys.argv[1])
        data = payload.get("data")
        template_path = payload.get("template_path")
        output_path = payload.get("output_path")
        result = generate_test_plan(data, template_path, output_path)
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
