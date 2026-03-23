"""
AI Job Assistant - Main Script
================================
A complete AI-powered tool to analyze resumes, score them against JDs,
generate updated resumes, and produce cover letters.

Naming Convention (Updated Resumes):
    Harshad Shinde_QA_resume_[COMPANY_INITIALS].pdf
    e.g.  Harshad Shinde_QA_resume_PW.pdf  (Physics Wala)

Usage:
    python job_assistant.py

Requirements:
    pip install anthropic python-docx PyPDF2 reportlab
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path

# ─── Try importing dependencies ──────────────────────────────────────────────
try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False
    print("WARNING: anthropic not installed. Run: pip install anthropic")

try:
    import PyPDF2
    HAS_PDF_READ = True
except ImportError:
    HAS_PDF_READ = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("WARNING: reportlab not installed. Run: pip install reportlab")

# ─── Configuration ────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).parent.parent   # job-assistant-ai/
TEMPLATES_DIR = BASE_DIR / "templates"
RESUMES_DIR   = BASE_DIR / "resumes"
JD_DIR        = BASE_DIR / "job-descriptions"
RESULTS_DIR   = BASE_DIR / "results"
SCORES_DIR    = RESULTS_DIR / "scores"
UPDATED_DIR   = RESULTS_DIR / "updated-resumes"
COVERS_DIR    = RESULTS_DIR / "cover-letters"

CANDIDATE_NAME = "Harshad Shinde"
RESUME_LABEL   = "QA_resume"


# ─── Utility: Company Initials ────────────────────────────────────────────────
def get_company_initials(company_name: str) -> str:
    """
    Extract uppercase initials from a company name.
    Examples:
        'Physics Wala'              -> 'PW'
        'Tata Consultancy Services' -> 'TCS'
        'Google'                    -> 'G'
        'infosys'                   -> 'I'
    """
    noise = {"pvt", "ltd", "llc", "inc", "co", "the", "and", "&"}
    words = re.split(r"[\s\-_\.]+", company_name.strip())
    initials = "".join(
        w[0].upper() for w in words
        if w.lower() not in noise and w
    )
    return initials if initials else company_name[:3].upper()


def resume_filename(company_initials: str) -> str:
    """Return the canonical resume PDF filename."""
    return f"{CANDIDATE_NAME}_{RESUME_LABEL}_{company_initials}.pdf"


# ─── File Reading Helpers ─────────────────────────────────────────────────────
def read_pdf(filepath: str) -> str:
    if not HAS_PDF_READ:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    text = []
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)


def read_docx(filepath: str) -> str:
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = Document(filepath)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def read_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return read_pdf(filepath)
    elif ext in [".docx", ".doc"]:
        return read_docx(filepath)
    else:
        return Path(filepath).read_text(encoding="utf-8")


# ─── LLM Call ─────────────────────────────────────────────────────────────────
def call_claude(system_prompt: str, user_message: str, api_key: str) -> str:
    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )
    return message.content[0].text


# ─── PDF Resume Builder (strict 2-page max) ───────────────────────────────────
def save_resume_as_pdf(resume_text: str, company_initials: str) -> Path:
    """
    Save the updated resume as a professionally formatted PDF.
    Strictly enforces a MAXIMUM of 2 pages (A4).
    Filename: Harshad Shinde_QA_resume_[INITIALS].pdf
    """
    if not HAS_REPORTLAB:
        raise ImportError("reportlab not installed. Run: pip install reportlab")

    out_path = UPDATED_DIR / resume_filename(company_initials)

    def build_pdf(font_size: float, leading: float, section_before: float):
        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=A4,
            leftMargin=1.5 * cm,
            rightMargin=1.5 * cm,
            topMargin=1.2 * cm,
            bottomMargin=1.2 * cm,
        )

        name_style = ParagraphStyle(
            "Name", fontName="Helvetica-Bold", fontSize=15,
            leading=19, textColor=colors.HexColor("#1a1a2e"),
            alignment=TA_CENTER, spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "Contact", fontName="Helvetica", fontSize=font_size - 0.5,
            leading=leading - 1, textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER, spaceAfter=3,
        )
        section_style = ParagraphStyle(
            "Section", fontName="Helvetica-Bold", fontSize=font_size,
            leading=leading, textColor=colors.HexColor("#1a1a2e"),
            spaceBefore=section_before, spaceAfter=2,
        )
        job_title_style = ParagraphStyle(
            "JobTitle", fontName="Helvetica-Bold", fontSize=font_size,
            leading=leading, textColor=colors.HexColor("#333333"), spaceAfter=1,
        )
        body_style = ParagraphStyle(
            "Body", fontName="Helvetica", fontSize=font_size,
            leading=leading, textColor=colors.HexColor("#222222"),
            spaceAfter=2, alignment=TA_JUSTIFY,
        )
        bullet_style = ParagraphStyle(
            "Bullet", fontName="Helvetica", fontSize=font_size,
            leading=leading, textColor=colors.HexColor("#222222"),
            spaceAfter=1, leftIndent=12,
        )
        updated_style = ParagraphStyle(
            "Updated", fontName="Helvetica-Oblique", fontSize=font_size - 0.3,
            leading=leading, textColor=colors.HexColor("#0057b8"),
            spaceAfter=1, leftIndent=12,
        )

        story = []
        story.append(Paragraph(CANDIDATE_NAME, name_style))

        in_header = True
        lines = resume_text.split("\n")

        for line in lines:
            raw = line.strip()
            if not raw:
                story.append(Spacer(1, 2))
                continue

            # Skip title lines that duplicate the candidate name
            if raw.upper() in ("UPDATED RESUME", "RESUME", CANDIDATE_NAME.upper()):
                continue

            is_updated = "[UPDATED]" in raw
            raw_clean = raw.replace("[UPDATED]", "").strip()

            # ALL CAPS lines = section headers
            if raw_clean.isupper() and len(raw_clean) > 3:
                story.append(HRFlowable(
                    width="100%", thickness=0.5,
                    color=colors.HexColor("#1a1a2e"), spaceAfter=2
                ))
                story.append(Paragraph(raw_clean, section_style))
                in_header = False
                continue

            # Bullet points
            if raw_clean.startswith(("•", "-", "*", "→", "▪")):
                text = raw_clean.lstrip("•-*→▪ ").strip()
                story.append(Paragraph(
                    f"• {text}",
                    updated_style if is_updated else bullet_style
                ))
                continue

            # Job title / company lines (contain |)
            if "|" in raw_clean or raw_clean.startswith("**"):
                clean = raw_clean.replace("**", "")
                story.append(Paragraph(
                    clean,
                    updated_style if is_updated else job_title_style
                ))
                continue

            # Contact / header block
            if in_header:
                story.append(Paragraph(raw_clean, contact_style))
                continue

            # Default body
            story.append(Paragraph(
                raw_clean,
                updated_style if is_updated else body_style
            ))

        doc.build(story)

    # First attempt at normal size
    build_pdf(font_size=8.5, leading=12.0, section_before=6)

    # Check page count; if > 2, shrink fonts and rebuild
    if HAS_PDF_READ:
        with open(str(out_path), "rb") as f:
            pages = len(PyPDF2.PdfReader(f).pages)
        if pages > 2:
            build_pdf(font_size=7.8, leading=10.5, section_before=4)

    return out_path


# ─── Save Helpers (scores / cover letters) ────────────────────────────────────
def save_text_result(content: str, folder: Path, company_initials: str, suffix: str) -> Path:
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"{suffix}_{company_initials}_{date_str}.txt"
    out_path = folder / filename
    out_path.write_text(content, encoding="utf-8")
    return out_path


def save_cover_letter_pdf(content: str, company_initials: str) -> Path:
    if not HAS_REPORTLAB:
        raise ImportError("reportlab not installed. Run: pip install reportlab")
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"{CANDIDATE_NAME}_CoverLetter_{company_initials}_{date_str}.pdf"
    out_path = COVERS_DIR / filename

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    body = ParagraphStyle(
        "CL", fontName="Helvetica", fontSize=10.5,
        leading=15, spaceAfter=8, alignment=TA_JUSTIFY,
    )
    story = []
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped:
            story.append(Paragraph(stripped, body))
        else:
            story.append(Spacer(1, 6))
    doc.build(story)
    return out_path


# ─── Core AI Functions ────────────────────────────────────────────────────────
def ats_score_only(resume_text: str, api_key: str) -> str:
    system = (
        "You are a professional ATS expert and resume reviewer with 15+ years in talent acquisition. "
        "Your analysis is precise, fair, and actionable."
    )
    user = f"""Analyze this resume and provide a detailed review:

1. Overall Result: [Score out of 10]
2. Effectivity: [Score out of 10] — How effectively skills and experiences are presented.
3. Layout and Design: [Score out of 10] — Visual appeal and organization.
4. Content Relevance: [Score out of 10] — Relevance and adequacy of information.
5. Grammar and Syntax: [Score out of 10] — Language quality and readability.
6. Impact: [Score out of 10] — How the resume stands out.

Use ✅ for positives and 🙈 for areas of improvement. Be specific and actionable.

RESUME:
{resume_text}"""
    return call_claude(system, user, api_key)


def jd_analysis(resume_text: str, jd_text: str, api_key: str) -> str:
    system = (
        "You are an analytical expert and professional recruiter with strong ATS knowledge. "
        "You provide thorough, data-supported analysis to help candidates optimize their resumes."
    )
    user = f"""Analyze the resume against the job description.

1. Overall Result: [Score out of 10]
2. Effectivity: [Score out of 10]
3. Layout and Design: [Score out of 10]
4. Content Relevance: [Score out of 10] — Relevance to this specific JD.
5. Grammar and Syntax: [Score out of 10]
6. Impact: [Score out of 10]

Use ✅ for positives and 🙈 for areas of improvement.

Then provide:
- Keyword Match Table: ALL JD keywords → ✅ Present / 🙈 Missing
- ATS Compatibility Verdict: STRONG / MODERATE / WEAK
- Top 5 Recommendations for this specific role

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}"""
    return call_claude(system, user, api_key)


def update_resume(resume_text: str, jd_text: str, api_key: str) -> str:
    system = (
        "You are a professional resume writer and ATS optimization expert. "
        "You tailor resumes authentically — never fabricating experience, only rephrasing "
        "and emphasizing what is genuinely there."
    )
    user = f"""Optimize this resume for the job description. Follow these steps exactly:

STEP 1 — Analyze the JD: Extract Must-Have skills, Nice-to-Have skills, and key keywords.
STEP 2 — Cross-Reference: Compare JD requirements vs resume. Mark ✅ present / 🙈 missing.
STEP 3 — Update: Rephrase bullet points with stronger action verbs and JD-aligned terminology.
          Mark every changed line with [UPDATED].
STEP 4 — Output the COMPLETE updated resume text, ready for PDF formatting.

⚠️  STRICT 2-PAGE RULE — The resume MUST fit within MAXIMUM 2 pages on A4:
    • Summary / Objective: maximum 3–4 lines
    • Each job entry: maximum 4–5 bullet points, each max 1–2 lines
    • Skills section: comma-separated inline, not one-per-line
    • Remove redundant, outdated, or low-impact content
    • Do NOT add new sections — only update existing ones
    • Do NOT fabricate experience

STEP 5 — After the resume, add a brief "Changes Summary" (what changed and why).

JOB DESCRIPTION:
{jd_text}

Original RESUME:
{resume_text}"""
    return call_claude(system, user, api_key)

def update_resume_docx(docx_path: str, jd_text: str, api_key: str, company_initials: str) -> Path:
    """
    Reads DOCX, extracts paragraphs > 20 chars, sends array to LLM,
    receives rewrites, applies them in-place, saves new DOCX.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx not installed.")
    
    from docx import Document
    doc = Document(docx_path)
    
    target_paragraphs = []
    original_texts = []
    
    def process_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text.strip()
            # Ignore very short lines, typical headers, and candidate name
            if len(text) > 20 and not text.isupper() and "|" not in text:
                target_paragraphs.append(p)
                original_texts.append(text)
                
    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                
    if not original_texts:
        print("  WARNING: No modifiable text found in DOCX. Return original.")
        return Path(docx_path)
        
    print(f"  Extracting {len(original_texts)} candidate sentences for optimization...")
    
    system = (
        "You are an ATS optimization expert. "
        "You will be given a JSON array of sentences/bullet points from a resume, "
        "and a job description. Your task is to rewrite these bullet points to better match "
        "the job description by incorporating its keywords and emphasizing relevant experience. "
        "Return ONLY a JSON array of strings, with the exact same length and order as the input array. "
        "If a bullet point does not need changing or is irrelevant to the JD, return it exactly as is."
    )
    
    user_payload = json.dumps(original_texts, indent=2)
    user = f"""JOB DESCRIPTION:
{jd_text}

RESUME BULLET POINTS (JSON ARRAY):
{user_payload}

Return ONLY a valid JSON array of strings containing the optimized bullet points. Do not include markdown formatting like ```json."""

    response_text = call_claude(system, user, api_key)
    
    try:
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        updated_texts = json.loads(response_text.strip())
        
        if len(updated_texts) != len(original_texts):
            print(f"  WARNING: LLM returned {len(updated_texts)} items, expected {len(original_texts)}. Using original.")
            return Path(docx_path)
            
        changes_made = 0
        for p, original, updated in zip(target_paragraphs, original_texts, updated_texts):
            if original != updated:
                p.text = p.text.replace(original, updated)
                changes_made += 1
                
        print(f"  Successfully tailored {changes_made} bullet points.")
    except Exception as e:
        print(f"  Failed to parse LLM response as JSON: {e}")
        return Path(docx_path)
        
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    out_filename = f"{CANDIDATE_NAME}_{RESUME_LABEL}_{company_initials}_{date_str}.docx"
    out_path = UPDATED_DIR / out_filename
    UPDATED_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    
    return out_path


def generate_cover_letter(
    resume_text: str, jd_text: str,
    company_name: str, position: str, api_key: str
) -> str:
    system = (
        "You are an expert professional writer specializing in career documents. "
        "You write compelling, authentic cover letters — tailored, specific, never generic."
    )
    user = f"""Write a professional cover letter for this application.

Candidate: {CANDIDATE_NAME}
Position: {position}
Company: {company_name}

Requirements:
- Professional greeting ("Dear Hiring Manager" if name unknown)
- Introduction: candidate name, exact position, genuine interest in role and company
- Experience paragraph: connect background directly to JD with specific examples
- Achievements: 1–2 quantifiable achievements relevant to the role
- Company alignment: show knowledge of company values/mission
- Strong closing with eagerness for interview and thank you
- MAX one page | No clichés | No buzzwords | Confident and engaging tone

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}"""
    return call_claude(system, user, api_key)


# ─── CLI Helpers ──────────────────────────────────────────────────────────────
def print_menu():
    print("\n" + "=" * 60)
    print("         AI JOB ASSISTANT  —  Main Menu")
    print("=" * 60)
    print("  1.  ATS Score Only        (resume, no JD needed)")
    print("  2.  JD Analysis           (resume vs job description)")
    print("  3.  Update Resume         (in-place DOCX edit)")
    print("  4.  Generate Cover Letter (tailored PDF)")
    print("  5.  Full Pipeline         (options 2 + 3 + 4 together)")
    print("  0.  Exit")
    print("=" * 60)

def list_and_pick_docx(folder: Path, label: str) -> str:
    files = sorted(folder.glob("*.docx"))
    if not files:
        return list_and_pick(folder, label)
    print(f"\n  Files in /{folder.name}/:")
    for i, f in enumerate(files, 1):
        print(f"    [{i}] {f.name}")
    choice = input(f"  Enter number or full path to {label}: ").strip()
    if choice.isdigit() and 1 <= int(choice) <= len(files):
        return str(files[int(choice) - 1])
    return input(f"  Full path to {label}: ").strip()

def list_and_pick(folder: Path, label: str) -> str:
    files = sorted(folder.glob("*.*"))
    if files:
        print(f"\n  Files in /{folder.name}/:")
        for i, f in enumerate(files, 1):
            print(f"    [{i}] {f.name}")
        choice = input(f"  Enter number or full path to {label}: ").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(files):
            return str(files[int(choice) - 1])
    return input(f"  Full path to {label}: ").strip()


# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    print("\n  Welcome to AI Job Assistant!")
    print(f"  Candidate : {CANDIDATE_NAME}")
    print("─" * 60)

    api_key = os.environ.get("ANTHROPIC_API_KEY") or input(
        "  Enter your Anthropic API key: "
    ).strip()
    if not api_key:
        print("  No API key provided. Exiting.")
        return

    company_full = input(
        "  Company name (full, e.g. 'Physics Wala'): "
    ).strip() or "Company"
    initials = get_company_initials(company_full)
    print(f"  Company initials : {initials}")
    print(f"  Resume filename  : {resume_filename(initials)}")

    while True:
        print_menu()
        choice = input("  Select option: ").strip()

        if choice == "0":
            print(f"\n  Good luck with your job search, {CANDIDATE_NAME}!\n")
            break

        # Load resume (needed for all options)
        resume_path = list_and_pick_docx(RESUMES_DIR, "your RESUME")
        try:
            resume_text = read_file(resume_path)
            print(f"  Resume loaded ({len(resume_text):,} chars)")
        except Exception as e:
            print(f"  Error reading resume: {e}")
            continue

        # ── 1. ATS Score Only ──────────────────────────────────────────────────
        if choice == "1":
            print("\n  Scoring resume ...")
            result = ats_score_only(resume_text, api_key)
            out = save_text_result(result, SCORES_DIR, initials, "ats_score")
            print(f"\n{result}")
            print(f"\n  Saved -> {out}")

        # ── 2. JD Analysis ─────────────────────────────────────────────────────
        elif choice == "2":
            jd_path = list_and_pick(JD_DIR, "JOB DESCRIPTION")
            jd_text = read_file(jd_path)
            print(f"  JD loaded ({len(jd_text):,} chars)")
            print("\n  Analysing resume against JD ...")
            result = jd_analysis(resume_text, jd_text, api_key)
            out = save_text_result(result, SCORES_DIR, initials, "jd_analysis")
            print(f"\n{result}")
            print(f"\n  Saved -> {out}")

        # ── 3. Update Resume → DOCX ─────────────────────────────────────────────
        elif choice == "3":
            jd_path = list_and_pick(JD_DIR, "JOB DESCRIPTION")
            jd_text = read_file(jd_path)
            print(f"  JD loaded ({len(jd_text):,} chars)")
            print("\n  Generating tailored resume (in-place DOCX via LLM) ...")
            out = update_resume_docx(resume_path, jd_text, api_key, initials)
            print(f"\n  Resume saved -> {out}")
            print(f"  Filename      : {out.name}")

        # ── 4. Cover Letter ────────────────────────────────────────────────────
        elif choice == "4":
            jd_path = list_and_pick(JD_DIR, "JOB DESCRIPTION")
            jd_text = read_file(jd_path)
            position = input("  Position applying for: ").strip()
            print("\n  Generating cover letter ...")
            cover_text = generate_cover_letter(
                resume_text, jd_text, company_full, position, api_key
            )
            out = save_cover_letter_pdf(cover_text, initials)
            print(f"\n{cover_text}")
            print(f"\n  Saved -> {out}")

        # ── 5. Full Pipeline ───────────────────────────────────────────────────
        elif choice == "5":
            jd_path = list_and_pick(JD_DIR, "JOB DESCRIPTION")
            jd_text = read_file(jd_path)
            position = input("  Position applying for: ").strip()

            print("\n  [1/3] JD Analysis ...")
            analysis = jd_analysis(resume_text, jd_text, api_key)
            save_text_result(analysis, SCORES_DIR, initials, "jd_analysis")
            print("  Done.")

            print("  [2/3] Updating resume -> DOCX (in-place) ...")
            resume_out = update_resume_docx(resume_path, jd_text, api_key, initials)
            print(f"  Done -> {resume_out.name}")

            print("  [3/3] Generating cover letter ...")
            cover_text = generate_cover_letter(
                resume_text, jd_text, company_full, position, api_key
            )
            cover_out = save_cover_letter_pdf(cover_text, initials)
            print(f"  Done -> {cover_out.name}")

            print("\n  Pipeline complete!")
            print(f"\n  Files saved in /results/:")
            print(f"    updated-resumes/ -> {resume_out.name}")
            print(f"    cover-letters/   -> {cover_out.name}")
            print(f"    scores/          -> jd_analysis_{initials}_*.txt")

        else:
            print("  Invalid option. Please try again.")


if __name__ == "__main__":
    main()